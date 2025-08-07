#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para converter template de tradução JSON para DOCX
"""

import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime

def add_page_break(doc):
    """Adiciona quebra de página"""
    doc.add_page_break()

def create_docx_from_json(json_file, output_file):
    """Cria documento DOCX a partir do JSON"""
    
    # Carregar dados JSON
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Criar documento
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título principal
    title = doc.add_heading('TEMPLATE DE TRADUÇÃO - APLICAÇÃO NOKEKOI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações do projeto
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run('Aplicação: ').bold = True
    info_para.add_run(data['metadata']['application'])
    
    info_para = doc.add_paragraph()
    info_para.add_run('Idioma Origem: ').bold = True
    info_para.add_run(data['metadata']['source_language'])
    
    info_para = doc.add_paragraph()
    info_para.add_run('Idioma Destino: ').bold = True
    info_para.add_run(data['metadata']['target_language'])
    
    info_para = doc.add_paragraph()
    info_para.add_run('Data de Extração: ').bold = True
    info_para.add_run(data['metadata']['extraction_date'])
    
    info_para = doc.add_paragraph()
    info_para.add_run('Versão: ').bold = True
    info_para.add_run(data['metadata']['version'])
    
    doc.add_paragraph()
    
    # Instruções para o tradutor
    instructions = doc.add_heading('INSTRUÇÕES PARA TRADUÇÃO', level=1)
    instructions.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('1. Traduza apenas o texto no campo "TRADUÇÃO"')
    doc.add_paragraph('2. Mantenha a formatação e estrutura original')
    doc.add_paragraph('3. Considere o contexto cultural indígena')
    doc.add_paragraph('4. Use termos técnicos apropriados quando necessário')
    doc.add_paragraph('5. Mantenha a consistência terminológica')
    
    doc.add_paragraph()
    
    # Processar cada categoria
    for category_name, category_data in data['texts_to_translate'].items():
        # Título da categoria
        category_title = doc.add_heading(category_data['description'].upper(), level=1)
        category_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Processar cada texto
        for text_item in category_data['texts']:
            # ID e Prioridade
            id_para = doc.add_paragraph()
            id_para.add_run(f"ID: {text_item['id']}").bold = True
            id_para.add_run(f" | Prioridade: {text_item['priority'].upper()}")
            
            # Texto original
            original_para = doc.add_paragraph()
            original_para.add_run("TEXTO ORIGINAL: ").bold = True
            original_para.add_run(text_item['original_text'])
            
            # Campo para tradução
            translation_para = doc.add_paragraph()
            translation_para.add_run("TRADUÇÃO: ").bold = True
            translation_para.add_run("_" * 50)  # Espaço para tradução
            
            # Contexto (se houver)
            if text_item.get('context'):
                context_para = doc.add_paragraph()
                context_para.add_run("CONTEXTO: ").bold = True
                context_para.add_run(text_item['context'])
            
            doc.add_paragraph()  # Espaço entre itens
        
        # Quebra de página entre categorias
        if category_name != list(data['texts_to_translate'].keys())[-1]:
            add_page_break(doc)
    
    # Salvar documento
    doc.save(output_file)
    print(f"✅ Documento DOCX salvo como '{output_file}'")

def main():
    """Função principal"""
    print("📄 Convertendo template de tradução para DOCX...")
    
    input_file = 'translation_template_updated.json'
    output_file = 'template_traducao_nokekoi.docx'
    
    try:
        create_docx_from_json(input_file, output_file)
        print("✅ Conversão concluída com sucesso!")
    except Exception as e:
        print(f"❌ Erro na conversão: {e}")

if __name__ == "__main__":
    main() 